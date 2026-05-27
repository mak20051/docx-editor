"""
Document Rich Text Editor — Flask + Quill.js
http://localhost:5123

Supported formats (read):  .docx .doc .odt .rtf .txt .html .htm
Supported formats (write): .docx (always) or .txt (if opened as .txt)

Data flow:
  file ──▶ file_to_delta() ──▶ Quill Delta JSON ──▶ editor
  editor ──▶ Quill Delta JSON ──▶ delta_to_docx() ──▶ .docx download
"""

import json
import os
import subprocess
import threading
import webbrowser
from pathlib import Path

from flask import Flask, jsonify, render_template_string, request, send_file
from docx import Document
from docx.shared import Pt

app = Flask(__name__)
PORT = int(os.environ.get("PORT", 5123))
IN_DOCKER = os.environ.get("DOCKER") == "1"

SUPPORTED = {".docx", ".doc", ".odt", ".rtf", ".txt", ".html", ".htm"}

state: dict = {
    "filepath": None,
    "orig_ext": ".docx",
    "delta": {"ops": [{"insert": "\n"}]},
}


# ── DOCX → Quill Delta ────────────────────────────────────────────────────────

def _para_to_ops(para) -> list:
    """Convert a single python-docx Paragraph to a list of Delta ops."""
    ops = []
    style = para.style.name if para.style else "Normal"

    for run in para.runs:
        if not run.text:
            continue
        attrs: dict = {}
        if run.bold is True:
            attrs["bold"] = True
        if run.italic is True:
            attrs["italic"] = True
        if run.underline is True:
            attrs["underline"] = True
        if run.font.strike is True:
            attrs["strike"] = True
        op: dict = {"insert": run.text}
        if attrs:
            op["attributes"] = attrs
        ops.append(op)

    line: dict = {}
    lname = style.lower()
    if "heading 1" in lname:
        line["header"] = 1
    elif "heading 2" in lname:
        line["header"] = 2
    elif "heading 3" in lname:
        line["header"] = 3
    elif lname in ("list bullet", "list bullet 2"):
        line["list"] = "bullet"
    elif lname in ("list number", "list number 2"):
        line["list"] = "ordered"
    elif lname == "block text":
        line["blockquote"] = True

    nl: dict = {"insert": "\n"}
    if line:
        nl["attributes"] = line
    ops.append(nl)
    return ops


def docx_to_delta(doc: Document) -> dict:
    """
    Convert python-docx Document to Quill Delta ops.
    Walks body in document order so table cells aren't skipped.

    body children: w:p (paragraph) | w:tbl (table) | w:sectPr (ignore)
    """
    from docx.oxml.ns import qn as _qn
    from docx.text.paragraph import Paragraph
    from docx.table import Table

    ops = []

    def walk_body(element):
        for child in element:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag == "p":
                para = Paragraph(child, doc)
                if para.text or para.runs:
                    ops.extend(_para_to_ops(para))
                else:
                    ops.append({"insert": "\n"})
            elif tag == "tbl":
                tbl = Table(child, doc)
                for row in tbl.rows:
                    row_texts = []
                    for cell in row.cells:
                        cell_text = " ".join(
                            p.text for p in cell.paragraphs if p.text
                        )
                        row_texts.append(cell_text)
                    line = " | ".join(row_texts)
                    if line.strip():
                        ops.append({"insert": line})
                    ops.append({"insert": "\n"})

    walk_body(doc.element.body)
    return {"ops": ops or [{"insert": "\n"}]}


# ── Quill Delta → DOCX ────────────────────────────────────────────────────────

def delta_to_docx(ops: list, path: str):
    """Convert Quill Delta ops to a .docx file."""
    doc = Document()
    doc.styles["Normal"].font.size = Pt(11)

    # Buffer of (text, inline_attrs) for the current paragraph
    buf: list[tuple[str, dict]] = []

    def flush(line_attrs: dict):
        header = line_attrs.get("header")
        lst    = line_attrs.get("list")
        quote  = line_attrs.get("blockquote")
        code   = line_attrs.get("code-block")

        if header == 1:
            para = doc.add_heading("", level=1)
        elif header == 2:
            para = doc.add_heading("", level=2)
        elif header == 3:
            para = doc.add_heading("", level=3)
        elif lst == "bullet":
            para = doc.add_paragraph(style="List Bullet")
        elif lst == "ordered":
            para = doc.add_paragraph(style="List Number")
        elif quote:
            para = doc.add_paragraph(style="Quote")
        else:
            para = doc.add_paragraph()

        for text, attrs in buf:
            run = para.add_run(text)
            if attrs.get("bold"):
                run.bold = True
            if attrs.get("italic"):
                run.italic = True
            if attrs.get("underline"):
                run.underline = True
            if attrs.get("strike"):
                run.font.strike = True
            if code or line_attrs.get("code-block"):
                run.font.name = "Courier New"
                run.font.size = Pt(10)

        buf.clear()

    INLINE = {"bold", "italic", "underline", "strike", "color", "background", "link", "code"}
    LINE   = {"header", "list", "blockquote", "code-block", "align", "indent", "direction"}

    for op in ops:
        insert = op.get("insert")
        if not isinstance(insert, str):
            continue
        attrs = op.get("attributes") or {}

        inline_attrs = {k: v for k, v in attrs.items() if k in INLINE}
        line_attrs   = {k: v for k, v in attrs.items() if k in LINE}

        parts = insert.split("\n")
        for i, part in enumerate(parts):
            if part:
                buf.append((part, inline_attrs))
            if i < len(parts) - 1:
                flush(line_attrs)

    if buf:
        flush({})

    doc.save(path)


# ── Multi-format reader ───────────────────────────────────────────────────────

def _text_to_delta(text: str) -> dict:
    ops = []
    for line in text.splitlines():
        if line:
            ops.append({"insert": line})
        ops.append({"insert": "\n"})
    return {"ops": ops or [{"insert": "\n"}]}


def file_to_delta(path: Path, ext: str) -> dict:
    """Read any supported file and return a Quill Delta."""
    ext = ext.lower()

    if ext == ".docx":
        return docx_to_delta(Document(path))

    if ext == ".doc":
        # Many .doc files are actually .docx with wrong extension — try that first
        try:
            return docx_to_delta(Document(path))
        except Exception:
            pass
        # True binary .doc — use antiword
        result = subprocess.run(
            ["antiword", "-w", "0", str(path)],
            capture_output=True, timeout=30,
        )
        if result.returncode != 0:
            err = result.stderr.decode(errors="replace").strip()
            raise RuntimeError(f"Could not open .doc file: {err or 'unsupported format'}")
        return _text_to_delta(result.stdout.decode("utf-8", errors="replace"))

    if ext == ".odt":
        from odf import text as odftext, teletype
        from odf.opendocument import load as odf_load
        doc = odf_load(str(path))
        paras = [teletype.extractText(p) for p in doc.getElementsByType(odftext.P)]
        return _text_to_delta("\n".join(paras))

    if ext == ".rtf":
        from striprtf.striprtf import rtf_to_text
        raw = path.read_text(errors="replace")
        return _text_to_delta(rtf_to_text(raw))

    if ext == ".txt":
        return _text_to_delta(path.read_text(errors="replace"))

    if ext in (".html", ".htm"):
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(path.read_text(errors="replace"), "html.parser")
        return _text_to_delta(soup.get_text("\n"))

    raise ValueError(f"Unsupported format: {ext}")


def delta_to_txt(ops: list) -> str:
    """Flatten Delta to plain text for .txt export."""
    out = []
    for op in ops:
        if isinstance(op.get("insert"), str):
            out.append(op["insert"])
    return "".join(out)


# ── HTML ──────────────────────────────────────────────────────────────────────

HTML = r"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>DOCX Editor</title>
<link href="https://cdn.quilljs.com/1.3.7/quill.snow.css" rel="stylesheet">
<style>
  /* ── theme tokens ── */
  html[data-theme="dark"] {
    --bg:        #1e1e2e;
    --bg-bar:    #181825;
    --bg-qbar:   #24273a;
    --bg-sub:    #11111b;
    --border:    #313244;
    --fg:        #cdd6f4;
    --fg-dim:    #6c7086;
    --fg-muted:  #45475a;
    --accent:    #89b4fa;
    --accent-fg: #1e1e2e;
    --quote-fg:  #a6adc8;
    --code-bg:   #11111b;
    --code-fg:   #a6e3a1;
    --link:      #89dceb;
    --dot:       #f38ba8;
    --btn-bg:    #313244;
    --toggle-lbl: "☀︎";
  }
  html[data-theme="light"] {
    --bg:        #f8f8f2;
    --bg-bar:    #ededf0;
    --bg-qbar:   #e0e0e8;
    --bg-sub:    #e4e4ed;
    --border:    #c8c8d4;
    --fg:        #1e1e2e;
    --fg-dim:    #6c6c80;
    --fg-muted:  #9090a0;
    --accent:    #1d6fce;
    --accent-fg: #ffffff;
    --quote-fg:  #44445a;
    --code-bg:   #e4e4ed;
    --code-fg:   #276b27;
    --link:      #0070c0;
    --dot:       #d7222a;
    --btn-bg:    #d4d4de;
    --toggle-lbl: "🌙";
  }

  * { box-sizing: border-box; margin: 0; padding: 0; }
  html, body {
    height: 100%;
    background: var(--bg);
    color: var(--fg);
    font-family: system-ui, sans-serif;
    transition: background .2s, color .2s;
  }

  /* ── file toolbar ── */
  #file-bar {
    background: var(--bg-bar);
    padding: 8px 14px;
    display: flex;
    align-items: center;
    gap: 6px;
    border-bottom: 1px solid var(--border);
    flex-shrink: 0;
    transition: background .2s;
  }
  #file-bar button {
    background: var(--btn-bg); color: var(--fg);
    border: none; border-radius: 6px;
    padding: 6px 14px; font-size: 13px; cursor: pointer;
    font-family: inherit; transition: background .15s, color .15s;
    white-space: nowrap;
  }
  #file-bar button:hover { background: var(--accent); color: var(--accent-fg); }

  /* theme toggle — right-aligned, icon-only */
  #theme-toggle {
    margin-left: auto;
    background: var(--btn-bg); color: var(--fg);
    border: none; border-radius: 6px;
    padding: 6px 11px; font-size: 15px; cursor: pointer;
    line-height: 1; transition: background .15s;
  }
  #theme-toggle:hover { background: var(--accent); color: var(--accent-fg); }

  #filepath {
    margin-left: 10px; color: var(--fg-dim); font-size: 12px;
    overflow: hidden; text-overflow: ellipsis; white-space: nowrap;
    max-width: 340px;
  }
  #modified-dot { color: var(--dot); font-size: 16px; display: none; }

  /* ── Quill overrides ── */
  #quill-wrapper {
    display: flex;
    flex-direction: column;
    height: calc(100vh - 82px);
  }

  .ql-toolbar.ql-snow {
    background: var(--bg-qbar);
    border: none;
    border-bottom: 1px solid var(--border);
    flex-shrink: 0;
    transition: background .2s;
  }
  .ql-toolbar .ql-stroke { stroke: var(--fg); transition: stroke .15s; }
  .ql-toolbar .ql-fill   { fill:   var(--fg); transition: fill .15s; }
  .ql-toolbar .ql-picker-label { color: var(--fg); transition: color .15s; }
  .ql-toolbar button:hover .ql-stroke,
  .ql-toolbar .ql-active .ql-stroke { stroke: var(--accent); }
  .ql-toolbar button:hover .ql-fill,
  .ql-toolbar .ql-active .ql-fill   { fill:   var(--accent); }
  .ql-toolbar .ql-picker-label:hover,
  .ql-toolbar .ql-active { color: var(--accent); }
  .ql-toolbar .ql-picker-options {
    background: var(--bg-qbar);
    border-color: var(--border);
    color: var(--fg);
  }

  .ql-container.ql-snow {
    border: none;
    flex: 1;
    overflow: hidden;
    background: var(--bg);
    font-size: 15px;
    transition: background .2s;
  }
  .ql-editor {
    color: var(--fg);
    padding: 32px 80px;
    line-height: 1.75;
    height: 100%;
    overflow-y: auto;
    transition: color .2s;
  }
  .ql-editor.ql-blank::before { color: var(--fg-muted); font-style: normal; }
  .ql-editor h1, .ql-editor h2, .ql-editor h3 { color: var(--accent); }
  .ql-editor blockquote {
    border-left: 4px solid var(--accent);
    color: var(--quote-fg);
    padding-left: 16px;
    margin-left: 0;
  }
  .ql-editor pre.ql-syntax {
    background: var(--code-bg);
    border-radius: 6px;
    color: var(--code-fg);
    font-size: 13px;
  }
  .ql-editor a { color: var(--link); }

  /* ── status bar ── */
  #status {
    background: var(--bg-sub);
    padding: 4px 16px;
    font-size: 11px; color: var(--accent);
    border-top: 1px solid var(--border);
    flex-shrink: 0;
    height: 24px;
    transition: background .2s;
  }

  #file-input { display: none; }
</style>
</head>
<body>

<div id="file-bar">
  <button onclick="triggerOpen()">Open…</button>
  <button onclick="newDoc()">New</button>
  <button onclick="saveDoc()">Save</button>
  <button onclick="saveAs()">Save As…</button>
  <span id="filepath">Untitled</span>
  <span id="modified-dot">●</span>
  <button id="theme-toggle" onclick="toggleTheme()" title="Toggle day / night"></button>
</div>

<div id="quill-wrapper">
  <div id="toolbar">
    <span class="ql-formats">
      <select class="ql-header">
        <option selected></option>
        <option value="1">H1</option>
        <option value="2">H2</option>
        <option value="3">H3</option>
      </select>
    </span>
    <span class="ql-formats">
      <button class="ql-bold"></button>
      <button class="ql-italic"></button>
      <button class="ql-underline"></button>
      <button class="ql-strike"></button>
    </span>
    <span class="ql-formats">
      <button class="ql-list" value="ordered"></button>
      <button class="ql-list" value="bullet"></button>
    </span>
    <span class="ql-formats">
      <button class="ql-blockquote"></button>
      <button class="ql-code-block"></button>
    </span>
    <span class="ql-formats">
      <button class="ql-link"></button>
    </span>
    <span class="ql-formats">
      <button class="ql-clean"></button>
    </span>
  </div>
  <div id="editor"></div>
</div>

<div id="status">Ready</div>
<input type="file" id="file-input" accept=".docx,.doc,.odt,.rtf,.txt,.html,.htm" onchange="openFile(this)">

<script src="https://cdn.quilljs.com/1.3.7/quill.min.js"></script>
<script>
const quill = new Quill('#editor', {
  theme: 'snow',
  modules: { toolbar: '#toolbar' },
  placeholder: 'Start typing or open a .docx file…',
});

const statusEl   = document.getElementById('status');
const filepathEl = document.getElementById('filepath');
const dot        = document.getElementById('modified-dot');
const toggleBtn  = document.getElementById('theme-toggle');
let modified = false;
let origExt  = '.docx';

// ── theme ─────────────────────────────────────────────────────────────────────
const ICONS = { dark: '☀︎', light: '🌙' };
function applyTheme(theme) {
  document.documentElement.setAttribute('data-theme', theme);
  toggleBtn.textContent = ICONS[theme];
  localStorage.setItem('theme', theme);
}
function toggleTheme() {
  applyTheme(document.documentElement.getAttribute('data-theme') === 'dark' ? 'light' : 'dark');
}
applyTheme(localStorage.getItem('theme') || 'dark');

// ── helpers ──────────────────────────────────────────────────────────────────
function setStatus(msg)    { statusEl.textContent = msg; }
function setFilepath(p)    { filepathEl.textContent = p ? p.split('/').pop() : 'Untitled'; filepathEl.title = p || ''; }
function setModified(v)    { modified = v; dot.style.display = v ? 'inline' : 'none'; }

quill.on('text-change', () => setModified(true));

// ── load initial state ────────────────────────────────────────────────────────
(async () => {
  const d = await (await fetch('/state')).json();
  quill.setContents(d.delta, 'silent');
  setFilepath(d.filepath);
  origExt = d.orig_ext || '.docx';
  setModified(false);
})();

// ── keyboard shortcuts ────────────────────────────────────────────────────────
document.addEventListener('keydown', e => {
  if (e.metaKey && !e.shiftKey && e.key === 's') { e.preventDefault(); saveDoc(); }
  if (e.metaKey &&  e.shiftKey && e.key === 's') { e.preventDefault(); saveAs(); }
  if (e.metaKey && e.key === 'o') { e.preventDefault(); triggerOpen(); }
  if (e.metaKey && e.key === 'n') { e.preventDefault(); newDoc(); }
});

// ── file commands ─────────────────────────────────────────────────────────────
function triggerOpen() { document.getElementById('file-input').click(); }

async function openFile(input) {
  if (!input.files.length) return;
  const fd = new FormData();
  fd.append('file', input.files[0]);
  input.value = '';
  const d = await (await fetch('/open', { method: 'POST', body: fd })).json();
  if (d.error) { setStatus('Error: ' + d.error); return; }
  quill.setContents(d.delta, 'silent');
  setFilepath(d.filepath);
  origExt = d.ext || '.docx';
  setModified(false);
  const saveNote = origExt === '.docx' ? '' : ' — saves as .docx';
  setStatus(`Opened ${origExt.toUpperCase()}  (${d.paragraphs} paragraphs)${saveNote}`);
}

function newDoc() {
  if (modified && !confirm('Discard unsaved changes?')) return;
  quill.setContents([{ insert: '\n' }], 'silent');
  setFilepath(null);
  setModified(false);
  fetch('/reset', { method: 'POST' });
  setStatus('New document');
}

async function saveDoc() {
  const delta = quill.getContents();
  const d = await (await fetch('/save', {
    method: 'POST',
    headers: { 'Content-Type': 'application/json' },
    body: JSON.stringify({ delta: delta }),
  })).json();
  if (d.error)       { setStatus('Save failed: ' + d.error); return; }
  if (d.needs_saveas) { saveAs(); return; }
  setModified(false);
  setFilepath(d.filepath);
  setStatus('Saved: ' + d.filepath);
}

async function saveAs() {
  const cur = filepathEl.textContent;
  const defaultName = cur !== 'Untitled'
    ? cur.replace(/\.(doc|odt|rtf|txt|html|htm)$/i, '.docx')
    : 'document.docx';
  const name = prompt('Filename (.docx or .txt):', defaultName);
  if (!name) return;
  const fname = (name.endsWith('.docx') || name.endsWith('.txt')) ? name : name + '.docx';
  const delta = quill.getContents();
  const resp = await fetch('/saveas', {
    method: 'POST',
    headers: { 'Content-Type': 'application/json' },
    body: JSON.stringify({ delta: delta, name: fname }),
  });
  if (!resp.ok) { setStatus('Save failed'); return; }
  const blob = await resp.blob();
  const url  = URL.createObjectURL(blob);
  const a    = document.createElement('a');
  a.href = url; a.download = fname; a.click();
  URL.revokeObjectURL(url);
  setModified(false);
  setFilepath(fname);
  setStatus('Downloaded: ' + fname);
}
</script>
</body>
</html>
"""


# ── Flask routes ──────────────────────────────────────────────────────────────

@app.route("/")
def index():
    return render_template_string(HTML)


@app.route("/state")
def get_state():
    return jsonify({"filepath": state["filepath"], "orig_ext": state["orig_ext"], "delta": state["delta"]})


@app.route("/open", methods=["POST"])
def open_docx():
    f = request.files.get("file")
    if not f:
        return jsonify({"error": "no file"})
    ext = Path(f.filename).suffix.lower()
    if ext not in SUPPORTED:
        return jsonify({"error": f"Unsupported format: {ext}"})
    try:
        tmp = Path("/tmp") / f.filename
        f.save(tmp)
        print(f"[open] {f.filename} ({tmp.stat().st_size} bytes)", flush=True)
        delta = file_to_delta(tmp, ext)
        op_count  = len(delta["ops"])
        para_count = len([op for op in delta["ops"] if op.get("insert") == "\n"])
        text_chars = sum(len(op["insert"]) for op in delta["ops"] if isinstance(op.get("insert"), str) and op["insert"] != "\n")
        print(f"[open] delta: {op_count} ops, {para_count} paras, {text_chars} text chars", flush=True)
        if text_chars == 0:
            print(f"[open] WARNING: extracted 0 text chars from {f.filename}", flush=True)
        state["filepath"] = str(tmp)
        state["orig_ext"] = ext
        state["delta"] = delta
        return jsonify({"delta": delta, "filepath": str(tmp), "paragraphs": para_count, "ext": ext})
    except Exception as exc:
        print(f"[open] ERROR: {exc}", flush=True)
        return jsonify({"error": str(exc)})


@app.route("/reset", methods=["POST"])
def reset():
    state["filepath"] = None
    state["orig_ext"] = ".docx"
    state["delta"] = {"ops": [{"insert": "\n"}]}
    return jsonify({"ok": True})


@app.route("/save", methods=["POST"])
def save():
    data = request.get_json()
    if not state["filepath"]:
        return jsonify({"needs_saveas": True})
    try:
        delta_to_docx(data["delta"]["ops"], state["filepath"])
        state["delta"] = data["delta"]
        return jsonify({"filepath": state["filepath"]})
    except Exception as exc:
        return jsonify({"error": str(exc)})


@app.route("/saveas", methods=["POST"])
def saveas():
    data = request.get_json()
    name = data.get("name", "document.docx")
    ops  = data["delta"]["ops"]

    # honour .txt export; everything else → .docx
    if name.endswith(".txt"):
        dest = str(Path("/tmp") / name)
        Path(dest).write_text(delta_to_txt(ops))
        mime = "text/plain"
    else:
        if not name.endswith(".docx"):
            name += ".docx"
        dest = str(Path("/tmp") / name)
        delta_to_docx(ops, dest)
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    try:
        state["filepath"] = dest
        state["delta"] = data["delta"]
        return send_file(dest, as_attachment=True, download_name=name, mimetype=mime)
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


# ── main ──────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    url = f"http://localhost:{PORT}"
    if not IN_DOCKER:
        threading.Timer(0.9, lambda: webbrowser.open(url)).start()
    print(f"DOCX Rich Text Editor → {url}")
    app.run(host="0.0.0.0", port=PORT, debug=False)
